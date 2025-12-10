from io import BytesIO
import pandas as pd
from docx import Document
from datetime import datetime

from .commercial_rfp_shared_logger import logger
from .commercial_rfp_config_loader import ConfigLoader


class DocLibraryCreator:
    def __init__(self):
        self.config_loader = ConfigLoader.get_instance()
        self.in_blob_client = self.config_loader.blob_service_client.get_container_client(
            self.config_loader.content_container_name
        )
        self.out_blob_client = self.config_loader.blob_service_client.get_container_client(
            self.config_loader.commercial_rfp_survey_content_doc_library
        )

    def read_excel_from_blob(self, blob_client, blob_name: str) -> pd.DataFrame:
        stream = blob_client.get_blob_client(blob_name).download_blob().readall()
        return pd.read_excel(BytesIO(stream), engine="openpyxl")

    def get_latest_rfp_content_library_blob(self, blob_client) -> str | None:
        prefix = "RFP_content_library_"
        suffix = ".xlsx"
        latest_name = None
        latest_dt = None

        for blob in blob_client.list_blobs():
            name = blob.name
            if not (name.startswith(prefix) and name.endswith(suffix)):
                continue
            try:
                ts_part = name[len(prefix): -len(suffix)]
                dt = datetime.strptime(ts_part, "%Y%m%d")
            except Exception:
                continue

            if latest_dt is None or dt > latest_dt:
                latest_dt = dt
                latest_name = name

        return latest_name

    def create_docx_content(self, main_file_name: str, row: pd.Series, response_col: str) -> bytes:
        doc = Document()
        doc.add_paragraph(f"Source File Name: {main_file_name}")

        fields = [
            ("Client Name", "client name"),
            ("RFP Type", "rfp type"),
            ("Consultant", "consultant"),
            ("Date", "date"),
            ("Question", "question"),
            (response_col.title(), response_col),
            ("SME", "sme"),
        ]

        for label, key in fields:
            if key in row.index:
                value = row.get(key, None)
                if pd.notna(value) and str(value).strip() != "":
                    doc.add_paragraph(f"{label}: {value}")

        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer.read()

    def commerercial_rfp_content_doc_library_creation(self):
        logger.info("Creating .docx files and uploading to blob storage container.")
        try:
            latest_blob_name = self.get_latest_rfp_content_library_blob(self.in_blob_client)
            if not latest_blob_name:
                logger.error(
                    "No valid RFP_content_library_{timestamp}.xlsx files found in the blob container."
                )
                return

            # Rebuild the content-doc library container from scratch each run.
            # SharePoint cleanup is handled separately in the citation mapper.
            for blob in self.out_blob_client.list_blobs():
                self.out_blob_client.get_blob_client(blob.name).delete_blob()

            # Read Excel
            df = self.read_excel_from_blob(self.in_blob_client, latest_blob_name)
            df.columns = df.columns.str.lower()

            # Determine which response column to use
            response_col = None
            for col in ("response", "fixed answer"):
                if col in df.columns:
                    response_col = col
                    break

            if response_col is None:
                logger.error(
                    "Neither 'response' nor 'fixed answer' column found in RFP content library file."
                )
                return

            has_key_hash = "key_hash" in df.columns

            for _, row in df.iterrows():
                if has_key_hash:
                    key_hash_val = row.get("key_hash")
                    if pd.isna(key_hash_val) or str(key_hash_val).strip() == "":
                        continue

                    base_name = str(key_hash_val).strip()
                    if base_name.lower().endswith(".docx"):
                        docx_file_name = base_name
                    else:
                        docx_file_name = f"{base_name}.docx"
                else:
                    # Fallback to old behavior: use the first column as a reference ID
                    reference_col = df.columns[0]
                    ref_val = row.get(reference_col, None)
                    if pd.isna(ref_val) or str(ref_val).strip() == "":
                        continue
                    if isinstance(ref_val, float) and ref_val.is_integer():
                        ref_val = int(ref_val)
                    docx_file_name = f"RFP_Content_Library_{ref_val}.docx"

                # Create .docx content with correct source file name
                docx_bytes = self.create_docx_content(latest_blob_name, row, response_col)
                self.out_blob_client.get_blob_client(docx_file_name).upload_blob(
                    docx_bytes,
                    overwrite=True,
                )

            logger.info("Docx creation and upload completed.")
        except Exception as e:
            logger.critical(f"Pipeline failed: {e}", exc_info=True)