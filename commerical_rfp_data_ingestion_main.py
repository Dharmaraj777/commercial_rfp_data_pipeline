# Commercial RFP Pipeline — Refactored Code (Modularized)

## commerical_rfp_data_ingestion_main.py

```python
import os
import json
from datetime import datetime
import pandas as pd

from commercial_rfp_shared_logger import logger
from commercial_rfp_data_ingestion_utils import UtilityFunctions, set_runtime_config
from commercial_rfp_oai_client import OaiClient
from commercial_rfp_output_manager import OutputManager


def main():
    today_date = datetime.now().strftime("%Y-%m-%d")
    log_file_name = f"commercial_rfp_pipeline_{today_date}.log"
    config_file = 'config.json'

    with open(config_file, 'r') as f:
        config_details = json.load(f)
        # Aligning with your example: pick the two key containers + logs here
        output_container_name = config_details['commercial_rfp_survey_content_doc_library']
        prompt_container_name = config_details['commercial_rfp_survey_content_library']
        survey_logs = config_details['commercial_rfp_logs']

    # Publish a single runtime config for downstream modules
    runtime_cfg = dict(config_details)
    runtime_cfg["today_date"] = today_date
    runtime_cfg["log_file_name"] = log_file_name
    set_runtime_config(runtime_cfg)

    logger.info("Starting Commercial RFP Execution Process.")
    try:
        oai_client = OaiClient()
        result_manager = OutputManager()
        utils_func = UtilityFunctions()

        # Keep the same signature/structure you use elsewhere
        utils_func.main_processing(
            oai_client,
            result_manager,
            output_container_name,
            prompt_container_name,
            'process_commercial_rfp_pipeline'
        )

        # Upload log after all files are processed
        utils_func.upload_log_to_blob(log_file_name, survey_logs)

    except Exception as e:
        logger.error(f"An error occurred in the main process: {e}")
        print(f"An error occurred in the main process: {e}")
        # still upload logs on failure
        UtilityFunctions().upload_log_to_blob(log_file_name, survey_logs)

    logger.info("************* END of Processing File ******************")


if __name__ == "__main__":
    main()
```

## commercial_rfp_data_ingestion_utils.py

```python
"""
Utility functions for Commercial RFP data pipeline.
Consolidates common helpers used across ingestion, doc creation, and index reset modules.
Also hosts a small runtime-config store so *main* can load variables once and
all modules can reuse them without touching ConfigLoader again.
"""
from __future__ import annotations

import io
import json
import urllib.parse
from dataclasses import dataclass
from datetime import datetime
from io import BytesIO
from typing import List, Optional, Tuple, Union

import pandas as pd
import requests
from azure.storage.blob import BlobServiceClient
from msal import ConfidentialClientApplication

# Optional imports guarded where needed
try:
    from docx import Document  # Only used by create_docx_content
except Exception:  # pragma: no cover
    Document = None  # type: ignore

# -----------------------------
# Runtime config set/get (initialized in main)
# -----------------------------
_runtime_config: dict = {}

def set_runtime_config(cfg: dict) -> None:
    global _runtime_config
    _runtime_config = dict(cfg or {})

def get_runtime_config() -> dict:
    # If main hasn't set it yet, fall back to file config for backward-compat
    return _runtime_config or load_config()

# -----------------------------
# Config / Blob helpers
# -----------------------------

def load_config(config_file: str = 'config.json') -> dict:
    with open(config_file, 'r') as f:
        return json.load(f)

def init_blob_client(conn_str: str, container: str):
    blob_client = BlobServiceClient.from_connection_string(conn_str)
    return blob_client.get_container_client(container)

def connect_to_blob_storage_output_container(container_name: str, blob_service_client: BlobServiceClient):
    return blob_service_client.get_container_client(container_name)

def upload_df_as_excel_to_blob(connection_string: str, container_name: str, blob_name: str, df: pd.DataFrame, logger=None):
    try:
        output = BytesIO()
        with pd.ExcelWriter(output, engine="openpyxl") as writer:
            df.to_excel(writer, index=False)
        output.seek(0)

        blob_service_client = BlobServiceClient.from_connection_string(connection_string)
        blob_client = blob_service_client.get_blob_client(container=container_name, blob=blob_name)
        blob_client.upload_blob(output.getvalue(), overwrite=True)
        if logger:
            logger.info(f"Uploaded to Azure Blob Storage: {blob_name}")
    except Exception as e:
        if logger:
            logger.error(f"Failed to upload to Blob: {blob_name}, error: {e}")
        raise

def upload_blob_bytes(container_client, file_name: str, content: bytes):
    container_client.get_blob_client(file_name).upload_blob(content, overwrite=True)

def read_excel_from_blob(container_client, blob_name: str) -> pd.DataFrame:
    stream = container_client.get_blob_client(blob_name).download_blob().readall()
    return pd.read_excel(BytesIO(stream), engine="openpyxl")

def get_latest_rfp_content_library_blob(container_client) -> Optional[str]:
    from datetime import datetime
    files = []
    for blob in container_client.list_blobs():
        name = blob.name
        if name.startswith("RFP_content_library_") and name.endswith(".xlsx"):
            try:
                date_part = name.replace("RFP_content_library_", "").replace(".xlsx", "")
                file_date = datetime.strptime(date_part, "%Y%m%d")
                files.append((name, file_date))
            except Exception:
                continue
    if not files:
        return None
    return max(files, key=lambda x: x[1])[0]

def delete_all_blobs_in_container(container_client):
    for blob in container_client.list_blobs():
        container_client.get_blob_client(blob.name).delete_blob()

def upload_log_to_blob(blob_name: str, cfg_or_loader: Union[dict, object], log_stream, logger):
    """Accepts either the runtime cfg dict or the old ConfigLoader instance."""
    try:
        if isinstance(cfg_or_loader, dict):
            connection_string = cfg_or_loader["storage_connection_string"]
            container_name = cfg_or_loader["commercial_rfp_logs"]
            blob_service_client = BlobServiceClient.from_connection_string(connection_string)
        else:
            # Backward compatibility with ConfigLoader
            blob_service_client = cfg_or_loader.blob_service_client
            container_name = cfg_or_loader.commercial_rfp_logs

        blob_client = blob_service_client.get_blob_client(container=container_name, blob=blob_name)
        log_data = log_stream.getvalue().encode('utf-8')
        blob_client.upload_blob(BytesIO(log_data), overwrite=True)
        if logger:
            logger.info(f"Log file {blob_name} uploaded successfully.")
    except Exception as e:
        if logger:
            logger.error(f"Failed to upload log file {blob_name}: {str(e)}")

# -----------------------------
# MS Graph helpers
# -----------------------------

def get_graph_access_token(cert_path: str, thumbprint: str, client_id: str, tenant_id: str) -> str:
    authority = f"https://login.microsoftonline.com/{tenant_id}"
    scope = ["https://graph.microsoft.com/.default"]
    with open(cert_path, "rb") as f:
        cert_bytes = f.read()
    app = ConfidentialClientApplication(
        client_id=client_id,
        authority=authority,
        client_credential={
            "private_key": cert_bytes,
            "thumbprint": thumbprint,
        }
    )
    result = app.acquire_token_for_client(scopes=scope)
    if "access_token" in result:
        return result["access_token"]
    raise Exception(f"MSAL Auth failed: {result.get('error')} - {result.get('error_description')}")

def download_latest_excel_from_sharepoint_folder(access_token: str, site_url: str, folder_path: str) -> Tuple[BytesIO, str]:
    headers = {"Authorization": f"Bearer {access_token}"}

    folder_path_plain = urllib.parse.unquote(folder_path).strip().strip("/")
    if not folder_path_plain:
        raise ValueError("folder_path cannot be empty. Example: 'AI Data Repository/1. Content Library'")

    parts = [p for p in folder_path_plain.split("/") if p]
    drive_name = parts[0]
    folder_rel = "/".join(parts[1:])

    # Resolve site ID
    parsed = urllib.parse.urlparse(site_url)
    hostname = parsed.hostname
    site_path = parsed.path.rstrip('/')
    site_resp = requests.get(f"https://graph.microsoft.com/v1.0/sites/{hostname}:{site_path}", headers=headers)
    site_resp.raise_for_status()
    site_id = site_resp.json()["id"]

    # Find drive by display name
    drives_resp = requests.get(f"https://graph.microsoft.com/v1.0/sites/{site_id}/drives", headers=headers)
    drives_resp.raise_for_status()
    drive_id = None
    for d in drives_resp.json().get("value", []):
        if d.get("name", "").strip().lower() == drive_name.strip().lower():
            drive_id = d["id"]
            break
    if not drive_id:
        raise ValueError(f"Drive (library) named '{drive_name}' not found on site {site_url}.")

    # List children
    if folder_rel:
        enc_folder = urllib.parse.quote(folder_rel.strip("/"), safe="/")
        list_url = f"https://graph.microsoft.com/v1.0/sites/{site_id}/drives/{drive_id}/root:/{enc_folder}:/children"
    else:
        list_url = f"https://graph.microsoft.com/v1.0/sites/{site_id}/drives/{drive_id}/root/children"
    items_resp = requests.get(list_url, headers=headers)
    items_resp.raise_for_status()
    items = items_resp.json().get("value", [])

    excel_exts = (".xlsx", ".xls", ".xlsm")
    files = [it for it in items if "file" in it and it.get("name", "").lower().endswith(excel_exts)]
    if not files:
        where_ = f"{drive_name}/{folder_rel}" if folder_rel else drive_name
        raise FileNotFoundError(f"No Excel files found under '{where_}'.")

    def _parse_iso_z(dt_str: str) -> datetime:
        return datetime.fromisoformat(dt_str.replace("Z", "+00:00"))

    latest = max(files, key=lambda x: _parse_iso_z(x["lastModifiedDateTime"]))
    latest_rel_path = f"{folder_rel}/{latest['name']}" if folder_rel else latest["name"]
    enc_item_path = urllib.parse.quote(latest_rel_path.strip("/"), safe="/")
    download_url = f"https://graph.microsoft.com/v1.0/sites/{site_id}/drives/{drive_id}/root:/{enc_item_path}:/content"
    dl_resp = requests.get(download_url, headers=headers)
    dl_resp.raise_for_status()
    return BytesIO(dl_resp.content), latest["name"]

# -----------------------------
# Pandas helpers
# -----------------------------

def convert_list_to_dataframe(data) -> pd.DataFrame:
    if data is not None and len(data) > 0:
        return pd.DataFrame(data[1:], columns=data[0])  # First row is header
    return pd.DataFrame()

def get_length(value):
    if isinstance(value, str):
        return len(value)
    elif isinstance(value, (list, tuple, set)):
        return len(value)
    elif isinstance(value, (int, float)):
        return 1
    else:
        return None

def parse_dates(date):
    if pd.isna(date):
        return pd.NaT
    for fmt in ('%m/%d/%Y', '%Y-%m-%d', '%d-%m-%Y'):
        try:
            return pd.to_datetime(str(date), format=fmt)
        except Exception:
            continue
    return pd.to_datetime(date, errors='coerce')

# -----------------------------
# Doc helpers (used by content library creation)
# -----------------------------

def detect_response_column(df: pd.DataFrame) -> str:
    for col in ("response", "fixed answer"):
        if col in df.columns:
            return col
    raise KeyError("Neither 'response' nor 'fixed answer' column found")

def create_docx_content(main_file_name: str, row: pd.Series, response_col: str) -> bytes:
    if Document is None:
        raise RuntimeError("python-docx is required but not installed.")
    doc = Document()
    doc.add_paragraph(f"Source File Name: {main_file_name}")

    fields = [
        ("Client Name", "client name"),
        ("RFP Type", "rfp type"),
        ("Consultant", "consultant"),
        ("Date", "date"),
        ("Question", "question"),
        (response_col.title(), response_col),
        ("SME", "sme"),
    ]
    for label, key in fields:
        value = row.get(key, None)
        if pd.notna(value) and value != "" and key in row:
            doc.add_paragraph(f"{label}: {value}")

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.read()

# -----------------------------
# Orchestration helpers to match your "main" structure
# -----------------------------
class UtilityFunctions:
    def main_processing(self, oai_client, result_manager, output_container_name: str, prompt_container_name: str, process_name: str):
        """
        Mirrors your survey engine entrypoint signature.
        For Commercial RFP, we run: (1) raw ingestion & cleaning, (2) doc library creation.
        """
        # Import locally to avoid circular imports
        from commercial_rfp_raw_data_ingestion_and_cleaning import DataIngestion
        from commercial_rfp_content_doc_library_creation import commerercial_rfp_content_doc_library_creation

        # Step 1: Clean & publish RFP content library Excel
        di = DataIngestion(ctx=get_runtime_config())
        di.commercial_rfp_data_cleaning()

        # Step 2: Build DOCX library from latest Excel
        commerercial_rfp_content_doc_library_creation()

    def upload_log_to_blob(self, log_file_name: str, logs_container_name: str) -> None:
        """Match your example: explicitly supply target logs container."""
        from commercial_rfp_shared_logger import log_stream, logger as _logger
        cfg = get_runtime_config().copy()
        # Force the specific logs container provided by caller
        cfg["commercial_rfp_logs"] = logs_container_name
        upload_log_to_blob(log_file_name, cfg, log_stream, _logger)
```

## commercial_rfp_raw_data_ingestion_and_cleaning.py

```python
# import azure.functions as func
import pandas as pd
import re
import numpy as np
from datetime import datetime
from io import BytesIO
import warnings

from commercial_rfp_shared_logger import logger, log_stream
from commercial_rfp_config_loader import ConfigLoader

# Utilities
from commercial_rfp_data_ingestion_utils import (
    load_config,
    get_graph_access_token,
    download_latest_excel_from_sharepoint_folder,
    upload_df_as_excel_to_blob,
    convert_list_to_dataframe,
    get_length as util_get_length,
    parse_dates as util_parse_dates,
    upload_log_to_blob,
)

warnings.filterwarnings("ignore")


class DataIngestion():
    # Keep the same helper method names but delegate to utils,
    # so function structure remains intact while logic is centralized.

    def load_config():
        return load_config()

    def __init__(self, ctx: dict | None = None):
        # Accept a context dict (initialized in main). Fallback to config file for backward compatibility.
        cfg = ctx or load_config()
        self.config_loader = ConfigLoader.get_instance()

        # Variables (initialized centrally in main; mirrored here)
        self.today_date = datetime.now().strftime("%Y-%m-%d")
        self.log_file_name = f"commercial_rfp_data_processor_logs_{self.today_date}.log"

        self.cert_path = cfg["sharepoint_cert_path"]
        self.thumbprint = cfg["sharepoint_cert_thumbprint"]
        self.client_id = cfg["sharepoint_client_id"]
        self.tenant_id = cfg["sharepoint_tenant_id"]

        self.sharepoint_site_url = cfg['commercial_rfp_sharepoint_site_url']
        self.input_folder_path = cfg['commercial_rfp_sharepoint_content_file_url']
        self.output_folder_path = cfg['commercial_rfp_sharepoint_content_output_folder_url']

        self.azure_connection_string = cfg['storage_connection_string']
        self.azure_output_container_name = cfg['commercial_rfp_survey_content_library']
        self.raw_container_name = cfg['commercial_rfp_survey_raw_data_files']

    # ---- Delegated helpers to keep structure ----
    def get_graph_access_token(self, cert_path, thumbprint, client_id, tenant_id):
        return get_graph_access_token(cert_path, thumbprint, client_id, tenant_id)

    def download_latest_excel_from_sharepoint_folder(self, access_token: str, site_url: str, folder_path: str) -> BytesIO:
        return download_latest_excel_from_sharepoint_folder(access_token, site_url, folder_path)

    def upload_to_blob(self, connection_string, container_name, blob_name, df):
        return upload_df_as_excel_to_blob(connection_string, container_name, blob_name, df, logger=logger)

    def convert_list_to_dataframe(self, data):
        return convert_list_to_dataframe(data)

    def get_length(self, value):
        return util_get_length(value)

    def parse_dates(self, date):
        return util_parse_dates(date)

    def upload_log_to_blob(self, blob_name, config_loader):
        return upload_log_to_blob(blob_name, config_loader, log_stream, logger)

    # ---- Core cleaning logic (unchanged structure) ----
    def clean_data(self, df):
        logger.info(f"Initial DataFrame shape: {df.shape}")
        df.columns = df.columns.str.lower()
        df = df.applymap(lambda x: re.sub(r'\s+', ' ', str(x)).strip())

        # Date column validation
        if 'date' not in df.columns:
            similar_cols = [col for col in df.columns if 'date' in col]
            if similar_cols:
                msg = f"'date' column not found. Did you mean: {similar_cols}?"
            else:
                msg = "'date' column not found in the data. Available columns: {}".format(list(df.columns))
            logger.critical(msg)
            raise KeyError(msg)

        # Parse and filter dates (last 36 months)
        df['date'] = df['date'].apply(self.parse_dates)
        logger.info(f"After parsing 'date', NaT count: {df['date'].isna().sum()}")
        df = df.dropna(subset=['date'])
        logger.info(f"After dropping NaT dates, shape: {df.shape}")

        cutoff_date = (pd.Timestamp.now() - pd.DateOffset(months=36)).date()
        df['date'] = pd.to_datetime(df['date']).dt.date
        df = df[df['date'] >= cutoff_date]
        logger.info(f"After filtering for last 36 months, shape: {df.shape}")

        # Required columns
        missing = [c for c in ['question', 'response'] if c not in df.columns]
        if missing:
            msg = f"Missing required columns: {missing}"
            logger.critical(msg)
            raise KeyError(msg)

        # Clean rows
        df = df[~df['question'].isna() & (df['question'].str.lower() != 'none')]
        logger.info(f"After removing rows where question is 'None', shape: {df.shape}")

        df = df[~df['response'].isna() & (df['response'].str.lower() != 'none')]
        logger.info(f"After removing rows where response is 'None', shape: {df.shape}")

        df = df[~df['response'].isna() & (df['response'].str.lower() != 'nan')]
        logger.info(f"After removing rows where response is 'Nan', shape: {df.shape}")

        df['value_length'] = df['response'].apply(self.get_length)
        df = df[df['value_length'] != 0]
        logger.info(f"After removing rows with empty response, shape: {df.shape}")

        df = df[~df['response'].str.lower().isin(['n/a', 'not applicable.'])]
        logger.info(f"After removing 'N/A' and 'Not applicable.', shape: {df.shape}")

        df = df[~df['question'].str.lower().isin(['contact'])]
        logger.info(f"After removing 'contact', shape: {df.shape}")

        df.reset_index(drop=True, inplace=True)
        return df

    def drop_duplicates_same_question_and_response(self, df):
        grouped_df = df.groupby(['question', 'response']).size().reset_index(name='count')
        duplicates = grouped_df[grouped_df['count'] > 1]
        logger.info(f"Total number of duplicates found (question & response are exact same): {duplicates['count'].sum()}")
        if not duplicates.empty:
            unique_duplicates_count = duplicates[['question', 'response']].drop_duplicates().shape[0]
            logger.info(f"Number of unique duplicate combinations: {unique_duplicates_count}")
            df_remove_dups = df.drop_duplicates(subset=['question', 'response'])
            logger.info(f"After removing duplicates (same question + same response): {df_remove_dups.shape}")
            return df_remove_dups
        else:
            return df

    def same_question_duplicate_response(self, df):
        if df is None:
            raise ValueError("The DataFrame is None. Please check the previous steps in your code.")
        duplicates = df[df.duplicated('question', keep=False)].sort_values(by=['question', 'date'])
        if not duplicates.empty:
            max_dates = duplicates.groupby('question')['date'].max().reset_index()
            df_kept = df[df['date'].isin(max_dates['date']) | ~df['question'].isin(duplicates['question'])]
        else:
            df_kept = df
        df_kept = df_kept.sort_values(by=['date', 'question'])
        logger.info(f"After removing duplicate questions with oldest dates: {df_kept.shape}")
        return df_kept

    def get_unique_date_question_with_longest_response(self, df):
        df_new = pd.DataFrame(df)
        df_new['char_count'] = df['response'].apply(len)
        df_sorted = df_new.loc[df_new.groupby(['question'])['char_count'].idxmax()]
        df_sorted = df_sorted.drop(columns=['char_count']).sort_values(by=['date', 'question'])
        logger.info(f"After removing duplicate questions with different responses by taking longest response: {df_sorted.shape}")
        return df_sorted

    def commercial_rfp_data_cleaning(self):
        logger.info('Processing commercial_rfp_data_cleaning request.')
        try:
            # Authenticate with MSAL/Graph
            access_token = self.get_graph_access_token(self.cert_path, self.thumbprint, self.client_id, self.tenant_id)
            # Download Excel from SharePoint
            input_stream, original_filename = self.download_latest_excel_from_sharepoint_folder(access_token, self.sharepoint_site_url, self.input_folder_path)
            df_rfp = pd.read_excel(input_stream, engine="openpyxl")
            df_rfp.columns = df_rfp.columns.str.lower()
            logger.info(f"Original dataset of 'RFP content': {df_rfp.shape}")

            # Upload raw file
            self.upload_to_blob(self.azure_connection_string, self.raw_container_name, original_filename, df_rfp)

            # Clean and de-duplicate
            clean_df = self.clean_data(df_rfp)
            filtered_df = self.drop_duplicates_same_question_and_response(clean_df)
            df_kept = self.same_question_duplicate_response(filtered_df)
            df_unique_date_question = self.get_unique_date_question_with_longest_response(df_kept)

            if 'value_length' in df_unique_date_question.columns:
                df_unique_date_question = df_unique_date_question.drop(columns=['value_length'])

            # Normalize response flags
            df_unique_date_question['response'] = df_unique_date_question['response'].str.replace(
                r'(?i)(CONFIRMED|CONFIRMED\.|Confirmed via BlueInsights\.|Confirmed via mail\.|Confirmed\.|Yes\.\s*Confirmed\.)',
                'Confirmed',
                regex=True
            )

            timestamp = datetime.now().strftime("%Y%m%d")
            rfp_filename = f"RFP_content_library_{timestamp}.xlsx"

            self.upload_to_blob(self.azure_connection_string, self.azure_output_container_name, rfp_filename, df_unique_date_question)

            logger.info("Commercial RFP data cleaning completed successfully.")
            self.upload_log_to_blob(self.log_file_name, self.config_loader)
        except Exception as e:
            logger.exception(f"Pipeline failed: {e}", exc_info=True)
            self.upload_log_to_blob(self.log_file_name, self.config_loader)
```

## commercial_rfp_content_doc_library_creation.py

```python
import pandas as pd
from datetime import datetime

from commercial_rfp_shared_logger import logger, log_stream

# Utils (pull runtime config set by main)
from commercial_rfp_data_ingestion_utils import (
    get_runtime_config,
    init_blob_client as util_init_blob_client,
    read_excel_from_blob as util_read_excel_from_blob,
    get_latest_rfp_content_library_blob as util_get_latest_rfp_content_library_blob,
    delete_all_blobs_in_container as util_delete_all_blobs_in_container,
    upload_blob_bytes,
    detect_response_column,
    create_docx_content,
    upload_log_to_blob,
)

# Preserve the function name/signature

def commerercial_rfp_content_doc_library_creation():
    logger.info('Creating .doc files and uploading to blob storage container.')
    cfg = get_runtime_config()
    try:
        # Azure Blob clients
        in_blob_client = util_init_blob_client(
            cfg["storage_connection_string"],
            cfg["commercial_rfp_survey_content_library"]
        )

        out_blob_client = util_init_blob_client(
            cfg["storage_connection_string"],
            cfg["commercial_rfp_survey_content_doc_library"]
        )

        latest_blob_name = util_get_latest_rfp_content_library_blob(in_blob_client)
        if not latest_blob_name:
            logger.error("No valid RFP_content_library_{timestamp}.xlsx files found in the blob container.")
            return

        util_delete_all_blobs_in_container(out_blob_client)

        # Read Excel
        df = util_read_excel_from_blob(in_blob_client, latest_blob_name)
        response_col = detect_response_column(df)
        reference_col = df.columns[0]

        for _, row in df.iterrows():
            ref_val = row.get(reference_col, None)
            if pd.isna(ref_val) or str(ref_val).strip() == "":
                continue
            if isinstance(ref_val, float) and ref_val.is_integer():
                ref_val = int(ref_val)
            docx_file_name = f"RFP_Content_Library_{ref_val}.docx"
            docx_bytes = create_docx_content(latest_blob_name, row, response_col)
            upload_blob_bytes(out_blob_client, docx_file_name, docx_bytes)

        # Upload logs using runtime cfg (no ConfigLoader here)
        today_date = cfg.get("today_date") or datetime.now().strftime("%Y-%m-%d")
        log_file_name = cfg.get("log_file_name") or f"commercial_rfp_data_processor_logs_{today_date}.log"
        upload_log_to_blob(log_file_name, cfg, log_stream, logger)
    except Exception as e:
        logger.critical(f"Pipeline failed: {e}", exc_info=True)
        today_date = cfg.get("today_date") or datetime.now().strftime("%Y-%m-%d")
        log_file_name = cfg.get("log_file_name") or f"commercial_rfp_data_processor_logs_{today_date}.log"
        upload_log_to_blob(log_file_name, cfg, log_stream, logger)

# Calling (optional; usually triggered from main)
# commerercial_rfp_content_doc_library_creation()
```

## commercial_rfp_delete_index_files_and_reset_indexer.py

```python
from datetime import datetime

from commercial_rfp_shared_logger import logger, log_stream

from azure.search.documents import SearchClient
from azure.core.credentials import AzureKeyCredential
from azure.search.documents.indexes import SearchIndexerClient

from commercial_rfp_data_ingestion_utils import get_runtime_config, upload_log_to_blob


def delete_all_documents_from_index(search_client):
    while True:
        results = list(search_client.search("*", select="chunk_id", top=1000))
        if not results:
            break
        batch = [{"@search.action": "delete", "chunk_id": doc["chunk_id"]} for doc in results]
        search_client.upload_documents(documents=batch)
        logger.info(f"Deleted batch with {len(batch)} documents.")


def reset_indexer(indexer_client, indexer_name):
    indexer_client.reset_indexer(indexer_name)
    logger.info(f"Indexer '{indexer_name}' reset successfully.")


def commercial_rfp_delete_indexed_files_and_reset_indexer():
    logger.info('Processing commercial_rfp_delete_indexed_files_and_reset_indexer request.')
    cfg = get_runtime_config()
    try:
        today_date = cfg.get("today_date") or datetime.now().strftime("%Y-%m-%d")
        log_file_name = cfg.get("log_file_name") or f"commercial_rfp_data_processor_logs_{today_date}.log"

        service_endpoint = cfg["cogsearch_endpoint"]
        api_key = cfg["cogsearch_api_key"]
        index = [cfg["commercial_rfp_survey_content_doc_library_index"]]
        indexer = [cfg.get("commercial_rfp_survey_indexer_name", "")]  # keep optional

        credential = AzureKeyCredential(api_key)
        indexer_client = SearchIndexerClient(endpoint=service_endpoint, credential=credential)

        for index_name in index:
            if not index_name:
                continue
            logger.info(f"Resetting index: {index_name}")
            search_client = SearchClient(endpoint=service_endpoint, index_name=index_name, credential=credential)
            delete_all_documents_from_index(search_client)
            logger.info(f"All documents deleted from '{index_name}'.")

        for indexer_name in indexer:
            if not indexer_name:
                continue
            logger.info(f"Resetting and running indexer: {indexer_name}")
            reset_indexer(indexer_client, indexer_name)

        upload_log_to_blob(log_file_name, cfg, log_stream, logger)

    except Exception as e:
        logger.exception("Error resetting Azure Search indexes.")
        today_date = cfg.get("today_date") or datetime.now().strftime("%Y-%m-%d")
        log_file_name = cfg.get("log_file_name") or f"commercial_rfp_data_processor_logs_{today_date}.log"
        upload_log_to_blob(log_file_name, cfg, log_stream, logger)

# calling is done from main
# commercial_rfp_delete_indexed_files_and_reset_indexer()
```

## commercial_rfp_shared_logger.py

```python
import logging
from io import StringIO
import sys

log_stream = StringIO()

logger = logging.getLogger("commercial_rfp_data_pipeline_logger")
logger.setLevel(logging.INFO)

# Avoid adding handlers multiple times
if not logger.handlers:
    stream_handler = logging.StreamHandler(log_stream)
    stream_handler.setLevel(logging.INFO)
    formatter = logging.Formatter('%(asctime)s - %(levelname)s - %(message)s')
    stream_handler.setFormatter(formatter)

    console_handler = logging.StreamHandler(sys.stdout)
    console_handler.setLevel(logging.INFO)
    console_handler.setFormatter(formatter)

    logger.addHandler(stream_handler)
    logger.addHandler(console_handler)
```

## commercial_rfp_config_loader.py

```python
import json
from azure.core.credentials import AzureKeyCredential
from azure.search.documents import SearchClient
from azure.storage.blob import BlobServiceClient

class ConfigLoader:
    _config_instance = None

    def __init__(self, config_file='config.json'):
        if not ConfigLoader._config_instance:
            # Load configuration from the config file only once
            with open(config_file, 'r') as f:
                self.config_details = json.load(f)
            
            self.connection_string = self.config_details['storage_connection_string']
            self.content_container_name = self.config_details['commercial_rfp_survey_content_library']
            self.prompt_container_name = self.config_details['commercial_rfp_survey_prompt_library']
            self.output_container_name = self.config_details['commercial_rfp_survey_ai_generated_output']
            self.commercial_rfp_logs = self.config_details['commercial_rfp_logs']
            self.citation_map_conatiner = self.config_details['commercial_rfp_survey_citation_map']
            self.files_status_container = self.config_details['commercial_rfp_survey_files_status']
            self.files_status_blob = self.config_details['commercial_rfp_files_processing_status']
            self.commercial_rfp_survey_raw_data_files = self.config_details['commercial_rfp_survey_raw_data_files']
            # Create a BlobServiceClient
            self.blob_service_client = BlobServiceClient.from_connection_string(self.connection_string)

            # Azure OpenAI Configuration
            self.openai_api_base = self.config_details['openai_api_base']
            self.openai_api_key = self.config_details['openai_api_key']
           
            self.embedding_model = self.config_details['openai_embedding_model']
    
            self.openai_api_version= self.config_details['large_model_api_version']
            self.deployment_id = self.config_details['large_model']
            
           
            # Azure Cognitive Search Configuration
            self.search_endpoint = self.config_details['cogsearch_endpoint']
            self.search_index_name = self.config_details['commercial_rfp_survey_content_doc_library_index']
            self.search_key = self.config_details['cogsearch_api_key']
            self.credential = AzureKeyCredential(self.search_key)
            self.search_client = SearchClient(self.search_endpoint, self.search_index_name, credential=self.credential)

            self.redis_host = self.config_details['azure_redis_host']
            self.redis_key = self.config_details['azure_redis_key']

            self.mapping_filename = self.config_details['commercial_rfp_mapping_filename']
            
            ConfigLoader._config_instance = self  # Cache the instance for later reuse

    @staticmethod
    def get_instance():
        if ConfigLoader._config_instance is None:
            ConfigLoader()  # Initialize if not already done
        return ConfigLoader._config_instance
```

## commercial_rfp_oai_client.py

```python
class OaiClient:
    """Placeholder to align with your main() structure.
    Extend this when you wire RFP prompts to Azure OpenAI for content generation.
    """
    def __init__(self):
        pass
```

## commercial_rfp_output_manager.py

```python
class OutputManager:
    """Placeholder output manager to align with your main() structure.
    Add blob writes / Excel packaging here as needed for downstream steps.
    """
    def __init__(self):
        pass
```
